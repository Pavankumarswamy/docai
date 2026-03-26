"""
change_tracker.py – Inline Change Tracking for DOCAI

Implements INLINE visual tracking directly in the output .docx file.
NO separate change report is generated.

Tracking conventions (high-contrast, human-reviewer friendly):
  - Deleted/replaced text : RED strikethrough + pale-red shading
  - Inserted/new text     : BOLD BLUE text + CYAN highlight background

The colour scheme is intentionally strong so reviewers can scan a document
and instantly distinguish untouched paragraphs from updated ones.

Colour markers used by html_exporter.py for browser preview:
  Deleted  → font.strike=True  AND font.color.rgb == DELETED_RED
  Inserted → font.bold=True    AND font.color.rgb == INSERTED_BLUE
"""

import logging
from typing import List

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from lxml import etree

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Colour constants — chosen for maximum contrast against plain document text
# ---------------------------------------------------------------------------
DELETED_RED    = RGBColor(0xC0, 0x00, 0x00)   # Strong red  — deleted text
INSERTED_BLUE  = RGBColor(0x00, 0x4E, 0xA6)   # Royal blue  — inserted text

# Legacy aliases kept for any external references
RED   = DELETED_RED
GREEN = INSERTED_BLUE   # renamed; kept for back-compat


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _add_highlight(run, colour_val: str = "cyan"):
    """
    Add a Word highlight colour to a run (XML level).
    Valid values: yellow, green, cyan, magenta, red, blue,
                  darkBlue, darkCyan, darkGreen, darkRed, darkYellow,
                  white, black, gray, darkGray.
    """
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn("w:highlight")):
        rPr.remove(old)
    hl = etree.SubElement(rPr, qn("w:highlight"))
    hl.set(qn("w:val"), colour_val)


def _add_shading(run, fill_hex: str):
    """
    Add an XML character-shading element to create a coloured background
    on a run (fills gaps highlight colours don't cover).
    fill_hex: 6-char hex without '#', e.g. 'FFE0E0'.
    """
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn("w:shd")):
        rPr.remove(old)
    shd = etree.SubElement(rPr, qn("w:shd"))
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)

# Legacy alias
def _add_highlight_shading(run):
    _add_highlight(run, "cyan")


def _build_deleted_run(para: Paragraph, text: str):
    """
    Deleted text: bold red + strikethrough + pale-red background shading.
    Visually: ~~old text~~ on a pink background.
    """
    run = para.add_run(text)
    run.font.strike      = True
    run.font.bold        = True
    run.font.color.rgb   = DELETED_RED
    _add_shading(run, "FFE0E0")   # pale red background
    return run


def _build_inserted_run(para: Paragraph, text: str):
    """
    Inserted text: bold royal-blue + cyan Word-highlight + light-blue shading.
    Visually: vivid blue text on cyan/light-blue background.
    """
    run = para.add_run(text)
    run.font.bold        = True
    run.font.color.rgb   = INSERTED_BLUE
    _add_highlight(run, "cyan")          # Word highlight band
    _add_shading(run, "D6EEFF")          # light-blue character shading
    return run


# ---------------------------------------------------------------------------
# Block helpers (same as edit_engine but kept self-contained)
# ---------------------------------------------------------------------------

def _get_blocks(doc: Document) -> list:
    blocks = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            blocks.append(Table(child, doc))
    return blocks


# ---------------------------------------------------------------------------
# Public track API
# ---------------------------------------------------------------------------

def apply_tracked_changes(doc: Document, edits: List[dict]) -> None:
    """
    Apply inline change tracking to a Document for all edits.

    For each edit:
      - Clears existing runs in the block
      - Writes: unchanged prefix + [strikethrough original] + [inserted new] + unchanged suffix

    Args:
        doc:   Mutated in-place.
        edits: List of edit dicts (must include original_text, new_text, id, type).
    """
    if not edits:
        return

    blocks = _get_blocks(doc)

    for edit in edits:
        b_id    = edit.get("id")
        e_type  = edit.get("type", "paragraph").lower()
        orig    = edit.get("original_text", "").strip()
        new_t   = edit.get("new_text", "").strip()

        if not new_t:
            continue

        if b_id is None or b_id >= len(blocks):
            continue

        block = blocks[b_id]

        if e_type == "paragraph" and isinstance(block, Paragraph):
            _track_paragraph(block, orig, new_t)

        elif e_type == "table" and isinstance(block, Table):
            r_idx = edit.get("row_index")
            c_idx = edit.get("col_index")
            if r_idx is not None and c_idx is not None:
                try:
                    cell = block.rows[r_idx].cells[c_idx]
                    # Use first paragraph of cell
                    if cell.paragraphs:
                        _track_paragraph(cell.paragraphs[0], orig, new_t)
                except IndexError:
                    pass


def _track_paragraph(para: Paragraph, original_text: str, new_text: str) -> None:
    """
    Replace paragraph runs with tracked-change runs.
    Pattern: [prefix][strikethrough original][green new][suffix]
    """
    full_text = para.text

    # Clear existing runs
    for run in para.runs:
        run.text = ""

    if original_text and original_text in full_text:
        idx = full_text.find(original_text)
        prefix = full_text[:idx]
        suffix = full_text[idx + len(original_text):]

        if prefix:
            para.add_run(prefix)

        if original_text.strip():
            _build_deleted_run(para, original_text)

        _build_inserted_run(para, new_text)

        if suffix:
            para.add_run(suffix)

    elif original_text:
        # Original not found — append new text after existing content
        para.add_run(full_text)
        _build_inserted_run(para, " " + new_text)

    else:
        # Full replacement (no original specified)
        _build_inserted_run(para, new_text)
