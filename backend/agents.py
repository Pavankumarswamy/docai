import os
import json
import logging
import shutil
from pathlib import Path
from datetime import datetime, timezone

from document_processor import (
    extract_problems_from_excel,
    extract_sections_from_docx,
    apply_edits_to_docx
)
from results_generator import generate_results
from multi_agents import multi_agent_pipeline
from docx import Document

logger = logging.getLogger(__name__)

def run_pipeline(
    run_id: str,
    doc_folder: str,
    excel_file: str,
    team_name: str,
    leader_name: str,
    branch_name: str,
    runs: dict,
) -> dict:
    """
    Autonomous Document Healing pipeline.
    repo_url here is actually the local path to the folder.
    """
    start_time = datetime.now(timezone.utc)
    live = runs[run_id]["live"]
    all_fixes = []
    
    def update_live(phase: str | None = None, message: str | None = None, append_terminal: str | None = None):
        if phase: live["phase"] = phase
        if message: live["message"] = message
        if append_terminal: live["terminal_output"] += (append_terminal + "\n")
        logger.info(f"[{run_id}] [{phase}] {message}")

    try:
        folder_path = Path(doc_folder)
        excel_path = Path(excel_file)
        if not folder_path.exists() or not folder_path.is_dir():
            update_live("error", "The provided folder path does not exist or is not a directory.")
            return {}
        if not excel_path.exists() or not excel_path.is_file():
            update_live("error", "The provided Excel file path does not exist or is not a file.")
            return {}

        update_live("discovery", "Scanning folder for Word files...")
        
        # 1. Scan for files
        docx_files = [f for f in folder_path.rglob("*.docx") if not any(part.startswith(".") for part in f.parts)]
        
        # filter out temporary office files starting with ~$
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]
        excel_files = [excel_path]
        
        live["files"] = [{"path": f.name, "type": "file"} for f in docx_files] + \
                          [{"path": "edits_log.json", "type": "file"}, {"path": "results.json", "type": "file"}]
            
        update_live("discovery", f"Found {len(docx_files)} Word doc(s). Using explicit Excel file: {excel_path.name}")
        
        if not docx_files:
            update_live("done", "No Word documents found in the selected folder.")
            return {}
            
        # 2. Extract problems from Excel
        update_live("execution", f"Extracting problems from {excel_path.name}...")
        
        problems = extract_problems_from_excel(excel_path)
        update_live(append_terminal=f">>> Found {len(problems)} problem(s) in Excel.")
        
        if not problems:
            update_live("done", "No problems found in Excel.")
            return {}
            
        # Create CSV string representation
        csv_string = "### ALL CSV ROWS ###\n"
        for i, row in enumerate(problems):
            csv_string += f"Row {i+1}:\n" + "\n".join([f"  {k}: {v}" for k, v in row.items() if str(v).strip()]) + "\n\n"
            
        ci_timeline = []
        backup_dir = folder_path / ".backup"
        backup_dir.mkdir(exist_ok=True)
        
        final_status = "PASSED"
        
        for doc_idx, doc_file in enumerate(docx_files):
            doc_rel_path = str(doc_file.relative_to(folder_path)).replace("\\", "/")
            update_live("execution", f"Processing {doc_file.name} [{doc_idx+1}/{len(docx_files)}]")
            
            # Backup original
            backup_path = backup_dir / doc_file.name
            if not backup_path.exists():
                shutil.copy2(doc_file, backup_path)
                
            sections = extract_sections_from_docx(str(doc_file))
            if not sections:
                update_live(append_terminal=f">>> Failed to extract sections from {doc_file.name}")
                continue
            
            update_live(append_terminal=f">>> Detected {len(sections)} sections in {doc_file.name}")
                
            iter_start = datetime.now(timezone.utc)
            ci_timeline.append({
                "iteration": doc_idx + 1,
                "status": "PROCESS",
                "timestamp": iter_start.isoformat(),
                "problems_count": len(problems),
                "message": f"Processing {len(sections)} sections",
            })
            live["iterations"] = ci_timeline
            
            all_document_edits = []
            change_report_doc = Document()
            change_report_doc.add_heading(f"Change Tracking Report: {doc_file.name}", level=1)
            
            # Count elements
            sections_skipped = 0
            sections_processed = 0
            llm_calls_total = 0
            
            for s_idx, section in enumerate(sections):
                s_name = section["name"]
                
                if section["skip"]:
                    sections_skipped += 1
                    update_live(append_terminal=f"[SKIP] Ignored restricted or pre-intro section: {s_name}")
                    continue
                    
                update_live("fixing", f"Running Multi-Agent Pipeline on Section: {s_name}")
                sections_processed += 1
                llm_calls_total += 2 # Passes 1 and 2
                
                initial_state = {
                    "section_name": s_name,
                    "original_section": section["content"],
                    "csv_rows": csv_string,
                    "pass1_edits": [],
                    "final_edits": [],
                    "explanation": "",
                    "errors": []
                }
                
                try:
                    result_state = multi_agent_pipeline.invoke(initial_state)
                    
                    final_edits = result_state.get("final_edits", [])
                    explanation = result_state.get("explanation", "No explanation.")
                    
                    if not final_edits:
                        update_live(append_terminal=f"[{s_name}] No meaningful updates required -> Source retained.")
                        continue
                        
                    all_document_edits.extend(final_edits)
                    update_live(append_terminal=f"[{s_name}] Generated {len(final_edits)} edits. Reason: {explanation}")
                    
                    # Store fix info for json report
                    all_fixes.append({
                        "file": doc_rel_path,
                        "problem": f"Section Review: {s_name}",
                        "bug_type": "DOCUMENT_EDIT",
                        "error_message": explanation,
                        "status": "fixed",
                        "agent": "DOCAI-LangGraph",
                        "edits": final_edits
                    })
                    
                    # Append to change tracking word doc
                    heading = change_report_doc.add_heading(f"Section: {s_name}", level=2)
                    
                    change_report_doc.add_heading("Original Content Snippet (from context):", level=3)
                    orig_preview = section["content"][:500] + ("..." if len(section["content"]) > 500 else "")
                    change_report_doc.add_paragraph(orig_preview)
                    
                    change_report_doc.add_heading("Edits Applied (JSON):", level=3)
                    change_report_doc.add_paragraph(json.dumps(final_edits, indent=2))
                    
                    change_report_doc.add_heading("Summary of Changes:", level=3)
                    change_report_doc.add_paragraph(explanation)
                    
                    change_report_doc.add_heading("CSV Context Used:", level=3)
                    change_report_doc.add_paragraph("Full CSV problem queue passed to Multi-Agent Pipeline.")
                    
                    change_report_doc.add_paragraph("-" * 60)
                    
                except Exception as e:
                    update_live(append_terminal=f"[{s_name}] LLM Processing Error -> {e}. Retaining original content.")
                    continue
            
            date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")    
            
            if all_document_edits:
                update_live(append_terminal=f">>> Applying {len(all_document_edits)} total edits back to document...")
                output_name = f"{doc_file.stem}_{date_str}.docx"
                output_path = folder_path / output_name
                
                apply_edits_to_docx(str(doc_file), all_document_edits, str(output_path))
                update_live(append_terminal=f">>> Saved updated document: {output_name}")
            else:
                update_live(append_terminal=">>> No changes made to the document.")
                
            report_name = f"change_report_{doc_file.stem}_{date_str}.docx"
            report_path = folder_path / report_name
            change_report_doc.save(report_path)
            update_live(append_terminal=f">>> Saved change report: {report_name}")
            
            update_live(append_terminal=f"\n[SECTION STATS] Processed: {sections_processed} | Skipped: {sections_skipped} | LLM Calls: {llm_calls_total}\n")
                    
        # Write edits log
        edits_log_path = folder_path / "edits_log.json"
        with open(edits_log_path, "w", encoding="utf-8") as f:
            json.dump(all_fixes, f, indent=4)
            
        # Re-index files for frontend
        live["files"] = [{"path": str(f.relative_to(folder_path)).replace("\\", "/"), "type": "file"} for f in list(folder_path.rglob("*.*")) if not str(f.name).startswith("~$")]
            
        update_live("done", f"✅ Document processing complete. Log saved to edits_log.json")
        
        end_time = datetime.now(timezone.utc)
        results = generate_results(
            run_id=run_id,
            repo_url=doc_folder,
            team_name=team_name,
            leader_name=leader_name,
            branch_name="DOC_UPDATE",
            fixes=all_fixes,
            ci_iterations=ci_timeline,
            start_time=start_time,
            end_time=end_time,
            final_status="PASSED",
            output_dir=str(folder_path),
        )
        return results

    except Exception as e:
        logger.exception(f"Pipeline error: {e}")
        update_live("error", f"Pipeline failed: {e}")
        return {}


