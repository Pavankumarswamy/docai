import os
import io
import json
import logging
import mammoth
import pandas as pd
from pathlib import Path
from docx import Document
from datetime import datetime, timezone

from llm_client import _call_nvidia, _strip_markdown

logger = logging.getLogger(__name__)

def extract_problems_from_excel(excel_path: str) -> list[dict]:
    """Reads the first sheet of an Excel file and returns rows as dictionaries."""
    try:
        df = pd.read_excel(excel_path)
        # Drop rows where all elements are NaN
        df = df.dropna(how='all')
        # Convert to list of dicts, replacing NaNs with None/empty string
        df = df.fillna("")
        return df.to_dict('records')
    except Exception as e:
        logger.error(f"Failed to read Excel file {excel_path}: {e}")
        return []

def extract_text_and_tables_from_docx(doc_path: str) -> str:
    """Extracts structural text and tables from a Word document for the LLM to understand."""
    try:
        doc = Document(doc_path)
        content = []
        
        content.append(f"--- Document: {os.path.basename(doc_path)} ---")
        
        # We need to maintain some order. For simplicity, we just extract paragraphs then tables, 
        # or iterate over document elements. python-docx doesn't easily allow iterating both in order without internal body access.
        # So we'll iterate through blocks.
        
        from docx.document import Document as _Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph
        
        def iter_block_items(parent):
            if isinstance(parent, _Document):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                raise ValueError("something's not right")

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        for i, block in enumerate(iter_block_items(doc)):
            if isinstance(block, Paragraph):
                # Always include the paragraph ID even if empty to ensure 1:1 mapping with the blocks list
                text = block.text.strip()
                content.append(f"[Paragraph {i}] {text if text else '(Empty)'}")
            elif isinstance(block, Table):
                content.append(f"[Table {i}]")
                for r_idx, row in enumerate(block.rows):
                    row_data = []
                    for c_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.replace('\n', ' ').strip()
                        row_data.append(f"C{c_idx}: {cell_text if cell_text else ''}")
                    content.append(f"  Row {r_idx} | " + " | ".join(row_data))
                    
        return "\n".join(content)
    except Exception as e:
        logger.error(f"Failed to extract content from {doc_path}: {e}")
        return ""

def generate_document_fix(problem_desc: str, doc_content: str, run_id: str) -> dict:
    """Calls the LLM to figure out what edits to make based on the problem."""
    sys_prompt = (
        "You are an expert Document Editor AI. Your task is to analyze a problem description and the contents of a Word document and determine EXACTLY what text to replace.\n"
        "### RULES:\n"
        "1. Response MUST be a valid JSON object. No markdown blocks.\n"
        "2. PRESERVE STRUCTURE: Do NOT change section numbers (e.g. 4.2 to 4.3) or remove headers unless explicitly instructed by the bug report.\n"
        "3. PRECISION: Provide a specific 'original_text' fragment from the paragraph/cell. If your 'new_text' is different from 'original_text', a change is applied.\n"
        "4. ALIGNMENT: Use the [Paragraph X] or [Table X] ID provided in the context. These IDs must match your 'id' field exactly.\n"
        "{\n"
        '  "edits": [\n'
        '    {\n'
        '      "type": "paragraph" | "table",\n'
        '      "id": <index_number_from_context>,\n'
        '      "original_text": "The exact text to replace.",\n'
        '      "new_text": "The corrected text.",\n'
        '      "row_index": <int_or_null>,\n'
        '      "col_index": <int_or_null>\n'
        '    }\n'
        "  ],\n"
        '  "explanation": "Brief reasoning"\n'
        "}"
    )
    
    user_prompt = f"### Problem to Fix:\n{problem_desc}\n\n### Document Content:\n{doc_content}\n\nProvide the JSON with the required edits."
    
    messages = [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    try:
        response_text = _call_nvidia(messages)
        content = _strip_markdown(response_text)
        data = json.loads(content)
        return data
    except json.JSONDecodeError as e:
        logger.error(f"LLM JSON Decode error: {e}\nResponse: {response_text}")
        return {"edits": [], "explanation": f"Failed to parse LLM response: {str(e)}"}
    except Exception as e:
        logger.error(f"LLM Error: {e}")
        return {"edits": [], "explanation": f"LLM Call failed: {str(e)}"}

def apply_edits_to_docx(doc_path: str, edits: list, output_path: str) -> list:
    """Applies the parsed edits to the Word document and saves it cleanly (no markers)."""
    try:
        doc = Document(doc_path)
        applied_edits = []
        
        from docx.document import Document as _Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph
        
        def get_blocks(parent):
            blocks = []
            if isinstance(parent, _Document):
                parent_elm = parent.element.body
            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    blocks.append(Paragraph(child, parent))
                elif isinstance(child, CT_Tbl):
                    blocks.append(Table(child, parent))
            return blocks
            
        blocks = get_blocks(doc)
        
        for edit in edits:
            b_id = edit.get("id")
            if b_id is None or b_id >= len(blocks):
                continue
                
            block = blocks[b_id]
            e_type = edit.get("type")
            
            if e_type == "paragraph" and isinstance(block, Paragraph):
                orig = edit.get("original_text", "")
                new_t = edit.get("new_text", "")
                if orig and orig in block.text:
                    block.text = block.text.replace(orig, new_t)
                    applied_edits.append({**edit, "status": "applied"})
                elif not orig:
                    block.text = new_t # full replace
                    applied_edits.append({**edit, "status": "applied_full"})
                else:
                    applied_edits.append({**edit, "status": "failed_orig_not_found"})
                    
            elif e_type == "table" and isinstance(block, Table):
                r_idx = edit.get("row_index")
                c_idx = edit.get("col_index")
                orig = edit.get("original_text", "")
                new_t = edit.get("new_text", "")
                
                if r_idx is not None and c_idx is not None:
                    if r_idx < len(block.rows) and c_idx < len(block.rows[r_idx].cells):
                        cell = block.rows[r_idx].cells[c_idx]
                        if orig and orig in cell.text:
                            cell.text = cell.text.replace(orig, new_t)
                            applied_edits.append({**edit, "status": "applied"})
                        else:
                            cell.text = new_t
                            applied_edits.append({**edit, "status": "applied_cell_override"})
                    else:
                        applied_edits.append({**edit, "status": "failed_out_of_bounds"})
            else:
                applied_edits.append({**edit, "status": "failed_type_mismatch"})
                
        doc.save(output_path)
        return applied_edits
    except Exception as e:
        logger.error(f"Failed to apply edits to {doc_path}: {e}")
        return []

def docx_to_html(doc_path: str, highlight_type: str = None, highlight_edits: list = None) -> str:
    """Converts docx to HTML with dynamic, in-memory highlighting to keep source files clean."""
    try:
        if highlight_type and highlight_edits:
            doc = Document(doc_path)
            # Re-extract blocks for correct addressing
            from docx.document import Document as _Document
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.table import CT_Tbl
            from docx.table import _Cell, Table
            from docx.text.paragraph import Paragraph
            
            def get_blocks(parent):
                blocks = []
                if isinstance(parent, _Document):
                    parent_elm = parent.element.body
                else: # For _Cell, parent_elm is parent._tc
                    parent_elm = parent._tc
                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P): blocks.append(Paragraph(child, parent))
                    elif isinstance(child, CT_Tbl): blocks.append(Table(child, parent))
                return blocks
            
            blocks = get_blocks(doc)
            h_start = "≈FIX_S≈" if highlight_type == "after" else "≈OLD_S≈"
            h_end = "≈FIX_E≈" if highlight_type == "after" else "≈OLD_E≈"

            for edit in highlight_edits:
                b_id = edit.get("id")
                if b_id is None or b_id >= len(blocks): continue
                
                block = blocks[b_id]
                e_type = edit.get("type")
                target_text = edit.get("new_text" if highlight_type == "after" else "original_text", "")
                
                if not target_text: continue
                
                if e_type == "paragraph" and isinstance(block, Paragraph):
                    if target_text in block.text:
                        block.text = block.text.replace(target_text, f"{h_start}{target_text}{h_end}")
                    else:
                        # Full block highlight if text match fails but ID is correct
                        block.text = f"{h_start}{block.text}{h_end}"
                elif e_type == "table" and isinstance(block, Table):
                    r_idx, c_idx = edit.get("row_index"), edit.get("col_index")
                    if r_idx is not None and c_idx is not None:
                         if r_idx < len(block.rows) and c_idx < len(block.rows[r_idx].cells):
                             cell = block.rows[r_idx].cells[c_idx]
                             if target_text in cell.text:
                                 cell.text = cell.text.replace(target_text, f"{h_start}{target_text}{h_end}")
                             else:
                                 cell.text = f"{h_start}{cell.text}{h_end}"
            
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            result = mammoth.convert_to_html(doc_stream)
        else:
            with open(doc_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
        
        html = result.value
        # New highlights (Gold)
        h_tag_new = '<span style="background-color: #ffd700; color: #000; padding: 0 2px; border-radius: 2px; font-weight: bold; border-bottom: 1px solid #b8860b;">'
        html = html.replace("≈FIX_S≈", h_tag_new).replace("≈FIX_E≈", "</span>")
        
        # Old highlights (Light Coral / Reddish)
        h_tag_old = '<span style="background-color: #ffcccc; color: #b22222; padding: 0 2px; border-radius: 2px; text-decoration: line-through; opacity: 0.9;">'
        html = html.replace("≈OLD_S≈", h_tag_old).replace("≈OLD_E≈", "</span>")
        
        return html
    except Exception as e:
        logger.error(f"Mammoth conversion error for {doc_path}: {e}")
        return f"<div style='color:red;'>Failed to preview document: {str(e)}</div>"
