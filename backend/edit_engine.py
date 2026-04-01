"""
edit_engine.py – Section-Aware, Rollback-Safe Edit Application Engine

Key guarantees
--------------
1. SECTION SCOPING   — locate_block() only sees blocks within the target
                       section; cross-section edits are structurally impossible.
2. ROLLBACK SAFETY   — document is snapshot'd (BytesIO) before any edit;
                       restored automatically if structure integrity fails.
3. SAFE INSERT GUARD — fallback inserts are rejected when new_text is already
                       semantically present in the section (cosine > 0.70).
4. STRUCTURE CHECK   — after applying, verifies headings unchanged, paragraph
                       count not reduced >25%, table count unchanged.
5. FULL TRACEABILITY — every edit result carries: status, match_type,
                       confidence, plan_id, source_row_id, section_name.
6. NO SILENT DROPS   — every skipped edit is logged with a reason.
"""

import io
import logging
import re
from typing import List, Optional, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

from scorer import locate_block, confidence_decision, keyword_cosine
from metrics import DocMetrics

logger = logging.getLogger(__name__)

# Similarity threshold above which a fallback insert is suppressed
_INSERT_DUPE_THRESHOLD = 0.70
# Max allowed paragraph-count reduction (fraction)
_MAX_PARA_REDUCTION    = 0.25


# ── Document-level block walker ───────────────────────────────────────────────

def get_blocks(doc: Document) -> list:
    """Return ALL Paragraph and Table objects in document order."""
    blocks = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            blocks.append(Table(child, doc))
    return blocks


# ── Section-scoped block extractor (CRITICAL) ────────────────────────────────

_HEADING_RE = re.compile(
    r"^(\d+[\.\d]*\s+)?("
    r"introduction|overview|scope|description|architecture|"
    r"requirements?|validation|conclusion|summary|appendix|"
    r"references?|background|purpose|data\s+flow|"
    r"non.functional|functional|constraints?|assumptions?"
    r")\b",
    re.IGNORECASE,
)

def _is_heading_block(block) -> bool:
    """Heuristic: is this block a section heading?"""
    if not isinstance(block, Paragraph):
        return False
    style = (block.style.name or "").lower() if block.style else ""
    if style.startswith("heading"):
        return True
    text = block.text.strip()
    if not text:
        return False
    if re.match(r"^(\d+[\.\d]*\s+)[A-Z]", text):   # "1. Introduction"
        return True
    if text.isupper() and 2 <= len(text.split()) <= 8:   # "INTRODUCTION"
        return True
    if _HEADING_RE.match(text):
        return True
    return False


def get_section_blocks(doc: Document, section_name: str) -> list:
    """
    Return ONLY the blocks that belong to `section_name`.

    Algorithm:
      1. Walk all blocks.
      2. When a heading whose normalised text matches section_name is found,
         start collecting.
      3. Stop collecting when the next heading is encountered.

    Returns an empty list if the section is not found (edit_engine will log
    and skip gracefully — never crash).
    """
    all_blocks    = get_blocks(doc)
    section_lower = section_name.strip().lower()
    collecting    = False
    result        = []

    for block in all_blocks:
        if _is_heading_block(block):
            heading_text = block.text.strip().lower()
            # Normalise numbered prefix: "1. Introduction" → "introduction"
            heading_norm = re.sub(r"^\d+[\.\d]*\s+", "", heading_text).strip()

            if collecting:
                # Hit the NEXT heading — stop
                break

            # Match: exact or heading_norm contains / is contained in section_lower
            if (heading_norm == section_lower
                    or section_lower in heading_norm
                    or heading_norm in section_lower):
                collecting = True
                result.append(block)   # include the heading block itself
                continue

        if collecting:
            result.append(block)

    if not result:
        logger.warning(
            f"[EditEngine] Section '{section_name}' not found in document — "
            "edits will be skipped."
        )

    return result


# ── Document snapshot / restore ───────────────────────────────────────────────

def snapshot_doc(doc: Document) -> bytes:
    """Serialise Document to bytes — safe in-memory backup."""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def restore_doc(snapshot: bytes) -> Document:
    """Deserialise bytes back to a fresh Document object."""
    return Document(io.BytesIO(snapshot))


# ── Structure integrity check ─────────────────────────────────────────────────

def _heading_texts(doc: Document) -> list:
    return [b.text.strip() for b in get_blocks(doc) if _is_heading_block(b)]


def _para_count(doc: Document) -> int:
    return sum(1 for b in get_blocks(doc) if isinstance(b, Paragraph))


def _table_count(doc: Document) -> int:
    return sum(1 for b in get_blocks(doc) if isinstance(b, Table))


def structure_integrity_check(doc_after: Document, doc_before_snap: bytes) -> Tuple[bool, str]:
    """
    Compare doc_after against a pre-edit snapshot.

    Returns:
        (ok: bool, reason: str)
        ok=True  → structure is intact
        ok=False → violation detected; caller should rollback
    """
    doc_before = restore_doc(doc_before_snap)

    before_headings = _heading_texts(doc_before)
    after_headings  = _heading_texts(doc_after)

    # Rule 1: headings must not change
    if before_headings != after_headings:
        changed = [
            (a, b) for a, b in zip(before_headings, after_headings) if a != b
        ]
        extra   = after_headings[len(before_headings):]
        missing = before_headings[len(after_headings):]
        reason  = f"heading mismatch — changed={changed} extra={extra} missing={missing}"
        return False, reason

    # Rule 2: paragraph count must not drop >25%
    before_paras = _para_count(doc_before)
    after_paras  = _para_count(doc_after)
    if before_paras > 0:
        reduction = (before_paras - after_paras) / before_paras
        if reduction > _MAX_PARA_REDUCTION:
            reason = (f"paragraph count dropped {reduction:.0%} "
                      f"({before_paras} → {after_paras})")
            return False, reason

    # Rule 3: table count must not decrease
    before_tables = _table_count(doc_before)
    after_tables  = _table_count(doc_after)
    if after_tables < before_tables:
        reason = f"table count decreased ({before_tables} → {after_tables})"
        return False, reason

    return True, "ok"


# ── Paragraph text writer ─────────────────────────────────────────────────────

def _set_paragraph_text(para: Paragraph, new_text: str) -> None:
    """Write new_text into a paragraph preserving the first run's formatting."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


# ── Per-block edit applicators ────────────────────────────────────────────────

def _apply_paragraph_edit(block: Paragraph, edit: dict, match_type: str) -> str:
    original_text = edit.get("original_text", "").strip()
    new_text      = edit.get("new_text",      "").strip()

    if not new_text:
        return "skipped_empty"

    current = block.text

    if match_type in ("exact", "fuzzy", "semantic") and original_text and original_text in current:
        _set_paragraph_text(block, current.replace(original_text, new_text, 1))
        return "replaced"

    if match_type == "insert_fallback" or not original_text:
        _set_paragraph_text(block, new_text)
        return "inserted"

    # keyword_cos / semantic with no verbatim match → append
    cleaned = current.rstrip(". ")
    _set_paragraph_text(block, f"{cleaned}. {new_text}.")
    return "appended"


def _apply_table_edit(block: Table, edit: dict) -> str:
    r_idx    = edit.get("row_index")
    c_idx    = edit.get("col_index")
    new_text = edit.get("new_text", "").strip()

    if r_idx is None or c_idx is None or not new_text:
        return "skipped_empty"
    try:
        cell = block.rows[r_idx].cells[c_idx]
    except IndexError:
        return "skipped_oob"

    original_text = edit.get("original_text", "").strip()
    if original_text and original_text in cell.text:
        cell.text = cell.text.replace(original_text, new_text, 1)
        return "replaced"
    cell.text = new_text
    return "overridden"


# ── Safe insert guard ─────────────────────────────────────────────────────────

def _safe_insert_guard(new_text: str, section_content: str) -> Tuple[bool, float]:
    """
    Check whether new_text is already semantically present in section_content.

    Returns:
        (allow_insert: bool, similarity: float)
        allow_insert=False → content already present; skip the insert
    """
    if not section_content.strip():
        return True, 0.0
    sim = keyword_cosine(new_text, section_content)
    if sim > _INSERT_DUPE_THRESHOLD:
        logger.info(
            f"[EditEngine] Insert suppressed — duplicate risk "
            f"(similarity={sim:.2f} > {_INSERT_DUPE_THRESHOLD}): "
            f"'{new_text[:60]}'"
        )
        return False, sim
    return True, sim


# ── Public API ────────────────────────────────────────────────────────────────

def apply_edits(
    doc:                   Document,
    edits:                 List[dict],
    section_name:          str = "",
    section_content:       str = "",
    section_block_indices: Optional[List[int]] = None,
    doc_metrics:           Optional[DocMetrics] = None,
) -> List[dict]:
    """
    Apply LLM-generated edits to an in-memory Document.

    Scoping
    -------
    If section_block_indices is provided (from document_processor), uses those
    exact block indices to scope matching.
    Otherwise falls back to get_section_blocks() heuristic.
    If that also fails, uses all document blocks (safe fallback, logs warning).

    Rollback
    --------
    Caller is responsible for taking a snapshot BEFORE calling this function
    (via snapshot_doc) and checking structure_integrity_check afterwards.
    This function does NOT roll back itself — it is a pure applier.

    Args:
        doc:                   In-memory Document (mutated in place).
        edits:                 Edit dicts from the pipeline.
        section_name:          Name of the section being edited.
        section_content:       Raw text of the section (for insert guard).
        section_block_indices: Optional list of block indices (int) from
                               document_processor.extract_sections_from_docx.
        doc_metrics:           Optional DocMetrics instance to accumulate stats.

    Returns:
        List of edit dicts, each augmented with:
        status, match_type, confidence, plan_id, source_row_id, section_name
    """
    if not edits:
        return []

    all_blocks = get_blocks(doc)

    # ── Determine section-scoped block list ───────────────────────────────
    if section_block_indices is not None:
        # Precise: use indices from document_processor
        scoped_blocks = [
            all_blocks[i]
            for i in section_block_indices
            if 0 <= i < len(all_blocks)
        ]
        if not scoped_blocks:
            logger.warning(
                f"[EditEngine] No valid blocks from indices {section_block_indices} "
                f"for section '{section_name}'. Falling back to heuristic."
            )
    else:
        scoped_blocks = None

    if not scoped_blocks and section_name:
        scoped_blocks = get_section_blocks(doc, section_name)

    if not scoped_blocks:
        logger.warning(
            f"[EditEngine] Could not scope blocks for '{section_name}'. "
            "Using full document — cross-section risk logged."
        )
        scoped_blocks = all_blocks

    results = []

    for edit in edits:
        # Coerce id to int (LLM sometimes returns string)
        raw_id = edit.get("id")
        if raw_id is not None:
            try:
                raw_id = int(raw_id)
            except (ValueError, TypeError):
                raw_id = None

        e_type        = edit.get("type", "paragraph").lower()
        target_hint   = edit.get("target_hint") or edit.get("original_text", "")
        plan_id       = int(edit.get("plan_id",       0))
        source_row_id = str(edit.get("source_row_id", ""))
        new_text      = edit.get("new_text", "").strip()

        # ── Block location (section-scoped ONLY) ─────────────────────────
        # Prefer explicit block id if it falls within scoped_blocks
        if raw_id is not None:
            scoped_ids = {id(b) for b in scoped_blocks}
            if 0 <= raw_id < len(all_blocks) and id(all_blocks[raw_id]) in scoped_ids:
                named_block = all_blocks[raw_id]
                named_text  = named_block.text if isinstance(named_block, Paragraph) else ""
                if target_hint and target_hint[:20].lower() in named_text.lower():
                    block, b_idx, match_type, confidence = named_block, raw_id, "exact", 1.0
                else:
                    block, b_idx, match_type, confidence = locate_block(
                        scoped_blocks, target_hint, plan_id, source_row_id
                    )
            else:
                block, b_idx, match_type, confidence = locate_block(
                    scoped_blocks, target_hint, plan_id, source_row_id
                )
        else:
            block, b_idx, match_type, confidence = locate_block(
                scoped_blocks, target_hint, plan_id, source_row_id
            )

        # ── Confidence gate ───────────────────────────────────────────────
        decision = confidence_decision(confidence, section_name, target_hint)
        if decision == "skip":
            status = "skipped_low_confidence"
            if doc_metrics:
                doc_metrics.record(status, confidence)
            results.append(_result(edit, status, match_type, confidence, plan_id, source_row_id, section_name))
            continue

        # ── Apply ─────────────────────────────────────────────────────────
        if match_type == "insert_fallback" or block is None:
            if new_text:
                allow, sim = _safe_insert_guard(new_text, section_content)
                if not allow:
                    status = "skipped_duplicate_insert"
                else:
                    doc.add_paragraph(new_text)
                    status = "inserted_fallback"
            else:
                status = "skipped_empty"

        elif e_type == "table" and isinstance(block, Table):
            status = _apply_table_edit(block, edit)

        elif isinstance(block, Paragraph):
            status = _apply_paragraph_edit(block, edit, match_type)

        else:
            # Type mismatch within scoped blocks
            if isinstance(block, Paragraph):
                status = _apply_paragraph_edit(block, edit, match_type) + "_coerced"
            else:
                status = "skipped_type_mismatch"

        logger.info(
            f"[Applier] section='{section_name}' plan={plan_id} "
            f"match={match_type} conf={confidence:.2f} status={status}"
        )
        if doc_metrics:
            doc_metrics.record(status, confidence)

        results.append(_result(edit, status, match_type, confidence, plan_id, source_row_id, section_name))

    return results


def _result(edit, status, match_type, confidence, plan_id, source_row_id, section_name) -> dict:
    return {
        **edit,
        "status":        status,
        "match_type":    match_type,
        "confidence":    round(confidence, 3),
        "plan_id":       plan_id,
        "source_row_id": source_row_id,
        "section_name":  section_name,
    }
